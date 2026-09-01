# -*- coding: utf-8 -*-
"""pandoc 生成的 docx → 中文学位论文样式后处理。

功能：
1. 插入封面页（标题 / 学院 / 专业 / 姓名 / 学号 / 指导教师占位 + 日期 + 分页符）；
2. 中文字体样式：正文宋体小四 1.5 倍行距、标题黑体、代码 Consolas；
3. 页边距（上 3.7 下 3.5 左 2.8 右 2.6 cm），首页不加页眉页脚；
4. 页眉（论文标题 + 下边框线）、页脚（第 X 页 共 Y 页，Word 域字段）；
5. 目录标题改为中文"目  录"，并在 settings 写入 updateFields 使 Word 打开时自动刷新目录。

注意事项：
- 不能用 doc.styles[key] 按名访问：python-docx 会经过 BabelFish 的 UI 名称转换，
  内置样式名称（如 "Heading 1"）转换后匹配失败。本脚本改为遍历 doc.styles 按
  st.name 匹配，最稳健。
- 设置显式字体时需移除 rFonts 上的 theme 属性（asciiTheme 等），否则 Word 仍可能
  使用主题字体；同理标题颜色需移除 themeColor/themeShade。

用法: python docx_postprocess.py <input.docx> <output.docx>
依赖: python-docx (pip install python-docx)
"""
import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

THESIS_TITLE_LINES = ["基于InsightFace与浏览器插件的", "B站视频明星人脸识别系统设计与实现"]
HEADER_TEXT = "基于InsightFace与浏览器插件的B站视频明星人脸识别系统设计与实现"
PLACEHOLDER_LINES = [
    "学院：＿＿＿＿＿＿＿＿＿＿",
    "专业：＿＿＿＿＿＿＿＿＿＿",
    "姓名：＿＿＿＿＿＿＿＿＿＿",
    "学号：＿＿＿＿＿＿＿＿＿＿",
    "指导教师：＿＿＿＿＿＿＿＿＿＿",
]


# ---------------------------------------------------------------- 工具函数

def find_style(doc, name):
    """按 UI 名称查找样式（遍历 doc.styles，避开 BabelFish 转换问题）。"""
    for st in doc.styles:
        if st.name == name:
            return st
    return None


def _clean_theme_fonts(rFonts):
    """移除 rFonts 上的主题字体属性，确保显式字体生效。"""
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        key = qn(attr)
        if rFonts.get(key) is not None:
            del rFonts.attrib[key]


def set_run_font(run, ascii_font, east_font, size_pt=None, bold=None):
    run.font.name = ascii_font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_font)
    _clean_theme_fonts(rFonts)


def style_font(style, ascii_font, east_font, size_pt=None, bold=None):
    style.font.name = ascii_font
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_font)
    _clean_theme_fonts(rFonts)


def set_style_color(style, hexval):
    """设置样式文字颜色并移除主题颜色属性。"""
    rPr = style.element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), hexval)
    for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
        key = qn(attr)
        if color.get(key) is not None:
            del color.attrib[key]


def add_text_run(paragraph, text, east="宋体", size=9):
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", east, size)
    return run


def add_field(paragraph, instr, size=9):
    """插入 Word 域（PAGE / NUMPAGES），打开文档时由 Word 计算。"""
    run = paragraph.add_run()
    set_run_font(run, "Times New Roman", "宋体", size)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    fld1.set(qn("w:dirty"), "true")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " " + instr + " "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(it)
    run._r.append(fld2)


# ---------------------------------------------------------------- 各处理步骤

def apply_styles(doc):
    normal = find_style(doc, "Normal")
    if normal is not None:
        style_font(normal, "Times New Roman", "宋体", 12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(0)

    for name, east, size, bold, center, before, after in [
        ("Heading 1", "黑体", 16, True, True, 18, 12),
        ("Heading 2", "黑体", 14, True, False, 12, 6),
        ("Heading 3", "黑体", 12, True, False, 6, 6),
    ]:
        st = find_style(doc, name)
        if st is None:
            continue
        style_font(st, "Times New Roman", east, size, bold)
        set_style_color(st, "000000")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.3
        if center:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name in ("First Paragraph", "Body Text"):
        st = find_style(doc, name)
        if st is not None:
            style_font(st, "Times New Roman", "宋体", 12)
            st.paragraph_format.line_spacing = 1.5
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    st = find_style(doc, "TOC Heading")
    if st is not None:
        style_font(st, "Times New Roman", "黑体", 16, True)
        set_style_color(st, "000000")
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    st = find_style(doc, "Source Code")
    if st is not None:
        style_font(st, "Consolas", "宋体", 9)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.space_after = Pt(0)

    st = find_style(doc, "Verbatim Char")
    if st is not None:
        style_font(st, "Consolas", "宋体", 10)


def rename_toc(doc):
    """把 pandoc 的 'Table of Contents' 改为中文 '目  录'。"""
    body = doc.element.body
    paras = body.xpath('.//w:p[w:pPr/w:pStyle[@w:val="TOCHeading"]]')
    for p in paras:
        for r in p.findall(qn("w:r")):
            p.remove(r)
        para = Paragraph(p, doc)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("目  录")
        set_run_font(run, "Times New Roman", "黑体", 16, bold=True)


def add_title_page(doc):
    """在目录域之前插入封面页，并以分页符收尾。"""
    body = doc.element.body
    sdt = body.find(qn("w:sdt"))
    created = []

    for _ in range(3):
        created.append(doc.add_paragraph())
    for line in THESIS_TITLE_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(line)
        set_run_font(run, "Times New Roman", "黑体", 22, bold=True)
        created.append(p)
    for _ in range(5):
        created.append(doc.add_paragraph())
    for label in PLACEHOLDER_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(label)
        set_run_font(run, "Times New Roman", "宋体", 14)
        created.append(p)
    for _ in range(3):
        created.append(doc.add_paragraph())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("二〇二六年八月")
    set_run_font(run, "Times New Roman", "宋体", 14)
    created.append(p)

    # 分页符段落（封面结束后另起一页进入目录）
    brk = doc.add_paragraph()
    brk.add_run().add_break(WD_BREAK.PAGE)

    # 把所有新建段落移动到目录域之前（保持顺序）
    for elem in [pp._p for pp in created] + [brk._p]:
        sdt.addprevious(elem)


def setup_header_footer(doc):
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.different_first_page_header_footer = True

    # 页眉：标题居中 + 下边框线
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(hp, HEADER_TEXT, size=9)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 页脚：第 X 页 共 Y 页
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(fp, "第 ", size=9)
    add_field(fp, "PAGE")
    add_text_run(fp, " 页 共 ", size=9)
    add_field(fp, "NUMPAGES")
    add_text_run(fp, " 页", size=9)


# ---------------------------------------------------------------- 主流程

def main():
    parser = argparse.ArgumentParser(description="中文学位论文 docx 后处理")
    parser.add_argument("input", help="pandoc 生成的 docx")
    parser.add_argument("output", help="输出 docx")
    args = parser.parse_args()

    doc = Document(args.input)
    apply_styles(doc)
    rename_toc(doc)
    add_title_page(doc)
    setup_header_footer(doc)

    # 目录域在 Word 打开时自动刷新
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)

    cp = doc.core_properties
    cp.title = THESIS_TITLE_LINES[0] + THESIS_TITLE_LINES[1]
    cp.author = "（待填写）"
    cp.subject = "毕业设计"

    doc.save(args.output)
    print("saved:", args.output)


if __name__ == "__main__":
    main()